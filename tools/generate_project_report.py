from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "CSE489_Chokro_Project_Report_Al_Sabah_Arnish.docx"
LOGO = ROOT / "assets" / "brand" / "chokro_app_icon.png"

# standard_business_brief preset with one named override:
# chokro_brand_accent_override = green headings/accent in place of preset blue.
GREEN = "087F5B"
DARK_GREEN = "075E4B"
MINT = "E8F5EF"
PALE_GREEN = "F3FAF7"
INK = "233238"
MUTED = "5D6B70"
LIGHT_BORDER = "C9DDD5"
WHITE = "FFFFFF"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, size=11, bold=None, italic=None, color=INK, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        element = tc_mar.find(qn(f"w:{margin}"))
        if element is None:
            element = OxmlElement(f"w:{margin}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=LIGHT_BORDER, size=6):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    for tag_name in ("tblW", "tblInd", "tblLayout"):
        existing = tbl_pr.find(qn(f"w:{tag_name}"))
        if existing is not None:
            tbl_pr.remove(existing)

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)

    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            existing = tc_pr.find(qn("w:tcW"))
            if existing is not None:
                tc_pr.remove(existing)
            tc_w = OxmlElement("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            tc_pr.insert(0, tc_w)
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_border(paragraph, color=GREEN, size=12, space=6):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relation_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), GREEN)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "22")
    run_pr.extend([fonts, color, underline, size])
    run.append(run_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_numbering_definition(doc, num_format, text, start=1):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
        if el.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
        if el.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start_el = OxmlElement("w:start")
    start_el.set(qn("w:val"), str(start))
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), num_format)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(fonts)
    level.extend([start_el, num_fmt, lvl_text, suff, lvl_jc, p_pr, r_pr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def set_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167


def add_list_item(doc, num_id, label, body):
    paragraph = doc.add_paragraph(style="List Text")
    set_num(paragraph, num_id)
    label_run = paragraph.add_run(label)
    set_run_font(label_run, bold=True, color=DARK_GREEN)
    body_run = paragraph.add_run(body)
    set_run_font(body_run)
    return paragraph


def add_bullet_link(doc, num_id, prefix, title, url, suffix=""):
    paragraph = doc.add_paragraph(style="List Text")
    set_num(paragraph, num_id)
    if prefix:
        prefix_run = paragraph.add_run(prefix)
        set_run_font(prefix_run, bold=True, color=DARK_GREEN)
    add_hyperlink(paragraph, title, url)
    if suffix:
        suffix_run = paragraph.add_run(suffix)
        set_run_font(suffix_run)
    return paragraph


def add_body(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph(style="Normal")
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True, color=DARK_GREEN)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_callout(doc, label, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.1
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), PALE_GREEN)
    shading.set(qn("w:val"), "clear")
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "6")
        border.set(qn("w:color"), LIGHT_BORDER)
        borders.append(border)
    p_pr.append(borders)
    label_run = paragraph.add_run(label)
    set_run_font(label_run, bold=True, color=DARK_GREEN)
    body_run = paragraph.add_run(text)
    set_run_font(body_run)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = rgb(DARK_GREEN)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.line_spacing = 1.0
    if title._element.pPr is not None:
        title_border = title._element.pPr.find(qn("w:pBdr"))
        if title_border is not None:
            title._element.pPr.remove(title_border)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(15)
    subtitle.font.color.rgb = rgb(MUTED)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(18)
    subtitle.paragraph_format.line_spacing = 1.1
    if subtitle._element.pPr is not None:
        subtitle_border = subtitle._element.pPr.find(qn("w:pBdr"))
        if subtitle_border is not None:
            subtitle._element.pPr.remove(subtitle_border)

    heading_tokens = {
        1: (16, 16, 8, GREEN),
        2: (13, 12, 6, GREEN),
        3: (12, 8, 4, DARK_GREEN),
    }
    for level, (size, before, after, color) in heading_tokens.items():
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    try:
        list_text = styles["List Text"]
    except KeyError:
        list_text = styles.add_style("List Text", WD_STYLE_TYPE.PARAGRAPH)
        list_text.base_style = normal
    list_text.font.name = "Calibri"
    list_text._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    list_text._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    list_text.font.size = Pt(11)
    list_text.font.color.rgb = rgb(INK)
    list_text.paragraph_format.space_before = Pt(0)
    list_text.paragraph_format.space_after = Pt(8)
    list_text.paragraph_format.line_spacing = 1.167


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    run = paragraph.add_run("CHOKRO PROJECT REPORT")
    set_run_font(run, size=9, bold=True, color=MUTED)
    run = paragraph.add_run("\tCSE 489")
    set_run_font(run, size=9, bold=True, color=MUTED)

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_before = Pt(0)
    footer_p.paragraph_format.space_after = Pt(0)
    run = footer_p.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    add_page_field(footer_p)

    first_footer = section.first_page_footer
    first_footer.paragraphs[0].text = ""


def set_core_properties(doc):
    props = doc.core_properties
    props.title = "Chokro Project Report"
    props.subject = "CSE 489 course project report"
    props.author = "Al Sabah Arnish"
    props.keywords = "Chokro, CSE 489, Flutter, Firebase, recycling, eco-rewards"
    props.comments = "Prepared from the implemented Chokro project workspace."


def add_cover(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)

    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p.paragraph_format.space_after = Pt(16)
    logo_run = logo_p.add_run()
    picture = logo_run.add_picture(str(LOGO), width=Inches(1.25))
    doc_pr = picture._inline.docPr
    doc_pr.set("name", "Chokro application icon")
    doc_pr.set("descr", "Green leaf icon representing the Chokro recycling application")

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(8)
    run = kicker.add_run("CSE 489  |  PROJECT REPORT")
    set_run_font(run, size=10, bold=True, color=GREEN)

    title = doc.add_paragraph("CHOKRO", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(
        "Verified Recycling and Eco-Rewards Platform", style="Subtitle"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    description = doc.add_paragraph()
    description.alignment = WD_ALIGN_PARAGRAPH.CENTER
    description.paragraph_format.space_after = Pt(26)
    run = description.add_run(
        "A cross-platform Flutter application for rewarding verified sustainable "
        "actions, supporting 3ZERO Greenpreneurs, and managing auditable eco-points."
    )
    set_run_font(run, size=11, italic=True, color=MUTED)

    metadata = doc.add_table(rows=5, cols=2)
    metadata.style = "Table Grid"
    set_table_geometry(metadata, [2160, 7200], indent_dxa=120)
    header_cells = metadata.rows[0].cells
    header_cells[0].text = "Field"
    header_cells[1].text = "Student Information"
    set_repeat_table_header(metadata.rows[0])
    for cell in header_cells:
        set_cell_shading(cell, GREEN)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, bold=True, color=WHITE)
    details = [
        ("Course", "CSE 489"),
        ("Name", "Al Sabah Arnish"),
        ("Student ID", "23101437"),
        ("Email", "________________________________________"),
    ]
    for row_index, (label, value) in enumerate(details, start=1):
        label_cell, value_cell = metadata.rows[row_index].cells
        set_cell_shading(label_cell, MINT)
        label_p = label_cell.paragraphs[0]
        label_p.paragraph_format.space_after = Pt(0)
        label_run = label_p.add_run(label)
        set_run_font(label_run, bold=True, color=DARK_GREEN)
        value_p = value_cell.paragraphs[0]
        value_p.paragraph_format.space_after = Pt(0)
        value_run = value_p.add_run(value)
        set_run_font(value_run)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_before = Pt(24)
    date_p.paragraph_format.space_after = Pt(0)
    date_run = date_p.add_run("23 August 2026")
    set_run_font(date_run, size=10, color=MUTED)
    date_p.add_run().add_break(WD_BREAK.PAGE)


def add_project_overview(doc):
    add_heading(doc, "1. Project Title", 1)
    add_callout(
        doc,
        "Project Title: ",
        "Chokro - Verified Recycling and Eco-Rewards Platform",
    )

    add_heading(doc, "2. Project Overview", 1)
    add_body(
        doc,
        "Chokro is a mobile-first recycling and eco-rewards application developed "
        "with Flutter and Firebase. It connects verified waste disposal, sustainable "
        "shopping, and donations through one auditable points wallet. The application "
        "targets Android and the web from a single Flutter codebase."
    )
    add_body(
        doc,
        "The system is designed around three user-facing profiles: 3ZERO Champion, "
        "3ZERO Greenpreneur, and 3ZERO Admin. Champions submit sustainable actions and "
        "use earned points; Greenpreneurs list eco-friendly products and manage orders; "
        "Admins review uncertain submissions, manage bins and users, configure policy, "
        "and observe platform statistics."
    )
    add_body(
        doc,
        "A trusted Node.js service performs sensitive verification and wallet changes. "
        "This keeps point awards, order settlement, and donation debits outside the "
        "client application and provides a stronger security boundary than direct "
        "client-side balance updates."
    )

    add_heading(doc, "3. Project Objectives", 1)
    objectives = add_numbering_definition(doc, "bullet", "•")
    add_list_item(
        doc,
        objectives,
        "",
        "Encourage correct waste disposal by converting verified sustainable actions into reward points.",
    )
    add_list_item(
        doc,
        objectives,
        "",
        "Create an accountable earn-and-spend cycle through a unified wallet and immutable transaction ledger.",
    )
    add_list_item(
        doc,
        objectives,
        "",
        "Give local eco-friendly entrepreneurs a dedicated marketplace and order-management channel.",
    )
    add_list_item(
        doc,
        objectives,
        "",
        "Provide administrators with review queues, policy controls, and meaningful platform statistics.",
    )


def add_features(doc):
    add_heading(doc, "4. Project Features", 1)
    add_body(
        doc,
        "By the end of this course, the following features have been implemented:"
    )
    letters = add_numbering_definition(doc, "lowerLetter", "%1)")
    features = [
        (
            "Login system. ",
            "Email/password registration, secure sign-in and sign-out, validation, "
            "account-state checks, and profile-aware routing are implemented with "
            "Firebase Authentication.",
        ),
        (
            "Reporting system. ",
            "The application provides administrative dashboards, live platform "
            "statistics, pending review queues, user submission histories, order "
            "records, wallet ledgers, and appeal outcomes for accountable monitoring.",
        ),
        (
            "Role and profile management. ",
            "3ZERO Admin, 3ZERO Greenpreneur, and 3ZERO Champion workspaces use a clear "
            "profile hierarchy. Switching profiles changes the workspace while the "
            "stored role remains the authority for protected actions.",
        ),
        (
            "Verified disposal workflow. ",
            "A four-step scan, photo, location, and confirmation process collects "
            "evidence. Server-side checks evaluate bin distance, photo provenance, "
            "duplicate hashes, screening results, and submission limits.",
        ),
        (
            "Eco-points wallet and transaction history. ",
            "Approved activity credits are recorded in an immutable ledger. Points can "
            "be spent during checkout or donated to selected green initiatives, with "
            "idempotent server operations to avoid duplicate charges or rewards.",
        ),
        (
            "Green marketplace and order system. ",
            "Greenpreneurs can publish and manage products. Champions can browse and "
            "filter the catalogue, maintain a cart, place multi-seller orders, apply "
            "points, follow order status, and confirm receipt.",
        ),
        (
            "Donation and payment prototypes. ",
            "The system supports point donations and clearly labelled bKash, Nagad, "
            "and card payment simulations without collecting credentials or moving "
            "real money.",
        ),
        (
            "Administrative review and appeals. ",
            "Flagged disposals, self-reported claims, Greenpreneur applications, and "
            "user appeals can be reviewed through dedicated admin queues with recorded "
            "decisions and responses.",
        ),
        (
            "Bin and notification management. ",
            "Admins can register bins and generate printable high-error-correction QR "
            "labels. Firebase Cloud Messaging supports decision and account-related "
            "notifications.",
        ),
        (
            "Security and validation. ",
            "Firestore rules enforce ownership and server-only fields, while the Node "
            "service performs trusted decisions, wallet mutations, rate limiting, and "
            "cross-record transaction checks.",
        ),
    ]
    for label, body in features:
        add_list_item(doc, letters, label, body)


def add_design(doc):
    add_heading(doc, "5. System Design and Technology", 1)
    add_body(
        doc,
        "Chokro separates presentation, state management, cloud data, and trusted "
        "business decisions. This architecture supports mobile and web interfaces while "
        "keeping sensitive balance and verification logic on the server."
    )

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "Layer"
    headers[1].text = "Implementation"
    rows = [
        ("Client application", "Flutter and Dart for Android and web from one codebase"),
        ("Navigation and state", "go_router and Riverpod for routing and reactive state"),
        ("Identity", "Firebase Authentication with email/password accounts"),
        ("Database", "Cloud Firestore with indexes and strict security rules"),
        ("Trusted service", "Node.js 22, Express, and Firebase Admin SDK"),
        ("Evidence capture", "QR scanning, camera/image handling, compression, and geolocation"),
        ("Media and screening", "Cloudinary-backed media handling and optional AI-assisted screening"),
        ("Notifications", "Firebase Cloud Messaging"),
        ("Quality assurance", "Flutter tests, Jest server tests, and Firestore emulator rules tests"),
    ]
    for layer, implementation in rows:
        cells = table.add_row().cells
        cells[0].text = layer
        cells[1].text = implementation
    set_table_geometry(table, [2520, 6840], indent_dxa=120)
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            if row_index == 0:
                set_cell_shading(cell, GREEN)
            elif row_index % 2 == 0:
                set_cell_shading(cell, PALE_GREEN)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.08
            for run in paragraph.runs:
                set_run_font(
                    run,
                    size=10.5,
                    bold=(row_index == 0 or col_index == 0),
                    color=WHITE if row_index == 0 else (DARK_GREEN if col_index == 0 else INK),
                )

    add_heading(doc, "5.1 Main Operational Workflow", 2)
    workflow = add_numbering_definition(doc, "decimal", "%1.")
    steps = [
        ("Account access. ", "A user registers or signs in and enters the correct profile workspace."),
        ("Action capture. ", "A Champion scans a registered bin, takes a photograph, confirms location, and submits evidence."),
        ("Trusted verification. ", "The server checks identity, location, photo evidence, duplicate risk, policy limits, and screening results."),
        ("Decision and reward. ", "A passing submission receives an atomic wallet credit; uncertainty is routed to human review."),
        ("Use of value. ", "The Champion spends points in the marketplace or donates them to a green initiative."),
        ("Monitoring. ", "Admins review queues, resolve appeals, manage policy and bins, and inspect aggregate statistics."),
    ]
    for label, body in steps:
        add_list_item(doc, workflow, label, body)


def add_resources(doc):
    add_heading(doc, "6. Online Resources Used", 1)
    add_body(
        doc,
        "The following online materials were used as learning references, implementation "
        "guides, and documentation sources during development."
    )

    add_heading(doc, "a) Reference", 2)
    refs = add_numbering_definition(doc, "bullet", "•")
    add_bullet_link(
        doc,
        refs,
        "W3Schools: ",
        "JavaScript JSON tutorial",
        "https://www.w3schools.com/js/js_json_intro.asp",
        " - background for JSON request and response structures.",
    )
    add_bullet_link(
        doc,
        refs,
        "Official documentation: ",
        "Flutter documentation",
        "https://docs.flutter.dev/",
        " - widgets, navigation, platform setup, and application development guidance.",
    )
    add_bullet_link(
        doc,
        refs,
        "Official documentation: ",
        "Add Firebase to a Flutter app",
        "https://firebase.google.com/docs/flutter/setup",
        " - FlutterFire and Firebase project configuration.",
    )
    add_bullet_link(
        doc,
        refs,
        "Official documentation: ",
        "Firebase email/password authentication for Flutter",
        "https://firebase.google.com/docs/auth/flutter/password-auth",
        " - account registration and login behavior.",
    )
    add_bullet_link(
        doc,
        refs,
        "YouTube video 1: ",
        "Firebase x Flutter Tutorial - Firebase Authentication",
        "https://www.youtube.com/watch?v=T96Pue6ePGA",
        " - practical email/password authentication walkthrough.",
    )
    add_bullet_link(
        doc,
        refs,
        "YouTube video 2: ",
        "Flutter Firebase Tutorial for Beginners",
        "https://www.youtube.com/watch?v=gEmFjbGYbRg",
        " - introduction to Firebase Authentication and Cloud Firestore with Flutter.",
    )

    add_heading(doc, "b) Stack Overflow or GitHub Links", 2)
    github = add_numbering_definition(doc, "bullet", "•")
    add_bullet_link(
        doc,
        github,
        "Project repository: ",
        "alsabaharnish/chokro",
        "https://github.com/alsabaharnish/chokro",
        " - source code and project history.",
    )
    add_bullet_link(
        doc,
        github,
        "GitHub reference: ",
        "firebase/flutterfire",
        "https://github.com/firebase/flutterfire",
        " - official Firebase plugins and examples for Flutter.",
    )
    add_bullet_link(
        doc,
        github,
        "GitHub reference: ",
        "firebase/firebase-admin-node",
        "https://github.com/firebase/firebase-admin-node",
        " - official Firebase Admin SDK for the trusted Node.js service.",
    )


def add_future_and_conclusion(doc):
    add_heading(doc, "7. Future Enhancements", 1)
    add_body(
        doc,
        "The following enhancements can be added to the current system to improve "
        "usability, reliability, reach, and operational value:"
    )
    letters = add_numbering_definition(doc, "lowerLetter", "%1)")
    enhancements = [
        (
            "Understanding of system. ",
            "Add a first-time guided tour, short role-based tutorials, contextual help, "
            "sample disposal demonstrations, and clearer explanations of how points are "
            "earned, spent, and donated.",
        ),
        (
            "Login system. ",
            "Add email verification, password reset, optional Google sign-in, multi-factor "
            "authentication for administrators, active-session management, and clearer "
            "recovery paths for locked or suspended accounts.",
        ),
        (
            "Reporting system. ",
            "Expand the dashboard with date and location filters, trend charts, bin-level "
            "utilization reports, Greenpreneur sales summaries, downloadable PDF/CSV "
            "exports, and scheduled administrator reports.",
        ),
        (
            "Production payments. ",
            "Replace payment simulations with approved bKash, Nagad, or card-gateway "
            "integrations using server-verified callbacks, secure receipts, reconciliation, "
            "refunds, and dispute handling.",
        ),
        (
            "Offline support, accessibility, and localization. ",
            "Allow evidence drafts to be captured with weak connectivity and synchronized "
            "later; add Bangla localization, larger-text support, screen-reader improvements, "
            "and simplified workflows for first-time smartphone users.",
        ),
        (
            "Smarter verification and production operations. ",
            "Improve duplicate detection, risk scoring, reviewer explanations, and model "
            "monitoring while adding centralized logs, automated backups, performance "
            "monitoring, stricter abuse controls, and an incident-response process.",
        ),
    ]
    for label, body in enhancements:
        add_list_item(doc, letters, label, body)

    add_heading(doc, "8. Conclusion", 1)
    add_body(
        doc,
        "Chokro demonstrates a complete sustainability reward cycle: users can verify an "
        "eco-friendly action, receive points through a trusted decision process, and reuse "
        "those points through marketplace purchases or initiative donations. The project "
        "also includes role-based administration, evidence review, appeals, reporting, and "
        "security controls appropriate for an academic prototype."
    )


def build_document():
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    set_core_properties(doc)
    add_cover(doc)
    add_project_overview(doc)
    add_features(doc)
    add_design(doc)
    add_resources(doc)
    add_future_and_conclusion(doc)

    # Ensure no widow-prone heading is left detached from its following content.
    for paragraph in doc.paragraphs:
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
